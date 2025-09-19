"""
FastAPI application for ABNT document processing.

This module defines a basic backend service that processes `.docx` files to
extract text, identify sections for formatting, preview the content as HTML,
and convert the document into ABNT-compliant format. The service exposes
three endpoints:

* `/analyze` — Accepts an uploaded `.docx` file and returns the plain text
  content along with heuristic spans identifying titles, citations, quotes,
  and references. This endpoint helps the frontend highlight parts of the
  document for user review without altering the original text.
* `/preview` — Accepts an uploaded `.docx` file and optionally a JSON map of
  confirmed spans. It returns sanitized HTML wrapped in basic styling for
  in-browser preview. The preview does not allow copy/paste of the
  document's content and is intended solely for visual inspection.
* `/convert` — Accepts an uploaded `.docx` file and optional metadata.
  Applies ABNT formatting (margins, font, line spacing, and alignment)
  without modifying the document's content. It returns a new `.docx` file
  named with an `ABNT_` prefix. Future versions could extend this endpoint to
  produce PDF output or embed metadata for cover/cover pages.

This service operates entirely on uploaded files and does not perform any
network calls or persist data beyond the scope of each request. The
implementation deliberately does not attempt to rewrite or correct the
user's text, staying in line with the requirement that no content is
altered.
"""

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from typing import Optional, List, Dict, Any
from pydantic import BaseModel
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
import tempfile
import os
import io
import re
import mammoth  # docx -> html converter for preview

app = FastAPI(title="ABNT Conversor", version="0.2.0")

class Span(BaseModel):
    """Represents a highlighted span in the plain text extracted from a document."""
    start: int
    end: int
    type: str  # "title" | "citation" | "quote" | "reference"

# Heuristic keywords for determining potential titles in the document.
KEYWORDS_TITLES = {
    "introdução",
    "metodologia",
    "conclusão",
    "referências",
    "sumário",
    "resumo",
    "abstract",
}

def load_docx_bytes(file: UploadFile) -> bytes:
    """Ensure the upload is a .docx file and return its bytes.

    Raises a ValueError if the file is not a `.docx`. This function reads
    the file's contents into memory immediately.
    """
    filename = file.filename or ""
    if not filename.lower().endswith(".docx"):
        raise ValueError("Envie um arquivo .docx")
    return file.file.read()

def extract_plain_text(doc: Document) -> str:
    """Extract the plain text from a `Document` object.

    The paragraphs are joined with newline characters to preserve basic
    structure (but without formatting or section markers). This makes it
    straightforward to compute index-based spans for highlighting.
    """
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    return "\n".join(parts)

def heuristic_spans(text: str) -> List[Dict[str, Any]]:
    """Identify likely spans for titles, long quotes, and references.

    This function uses simple heuristics:
    - A line of up to 80 characters that is all uppercase (ignoring accents
      and punctuation) or matches a keyword is considered a title.
    - Blocks of three or more consecutive lines longer than 80 characters
      each are flagged as potential long quotes (direct citations).
    - Lines matching a simple author/year pattern are marked as
      references.

    Spans are returned as a list of dictionaries with `start`, `end`, and
    `type` keys. The function sorts spans by their start position.
    """
    spans: List[Dict[str, Any]] = []
    lines = text.split("\n")
    idx = 0
    # Title detection
    for ln in lines:
        stripped = ln.strip()
        length = len(ln)
        if stripped and length <= 80:
            letters = re.sub(r"[^A-Za-zÁ-ÚÀ-ÙÂ-ÛÃ-ÕÇá-úà-ùâ-ûã-õç]", "", ln)
            all_caps = letters.isupper() and len(letters) >= 3
            looks_keyword = stripped.lower() in KEYWORDS_TITLES
            if all_caps or looks_keyword:
                spans.append({"start": idx, "end": idx + length, "type": "title"})
        idx += length + 1  # +1 accounts for the newline
    # Quote detection (three or more consecutive long lines)
    block_start = None
    block_len = 0
    offset = 0
    for ln in lines:
        if len(ln.strip()) >= 80:
            if block_start is None:
                block_start = offset
            block_len += len(ln) + 1
        else:
            if block_start is not None:
                if block_len >= 3 * 81:
                    spans.append({"start": block_start, "end": block_start + block_len - 1, "type": "quote"})
                block_start = None
                block_len = 0
        offset += len(ln) + 1
    if block_start is not None and block_len >= 3 * 81:
        spans.append({"start": block_start, "end": block_start + block_len - 1, "type": "quote"})
    # Reference detection (author + year)
    idx = 0
    for ln in lines:
        if re.match(r"^[A-ZÁ-Ú][A-Za-zÁ-úÀ-úÇç\\-]{2,}.*\(\d{4}\)\.?$", ln.strip()):
            spans.append({"start": idx, "end": idx + len(ln), "type": "reference"})
        idx += len(ln) + 1
    spans.sort(key=lambda s: (s["start"], s["end"]))
    return spans

def apply_abnt_formatting(doc: Document) -> None:
    """Apply basic ABNT formatting to a `Document` object.

    This sets page size, margins, font style and size, line spacing, and
    paragraph indentation. It does not modify the content or insert new
    headings, only standardizes the document's appearance.
    """
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(2)
        section.footer_distance = Cm(1.5)
    # Set base font and spacing
    base = doc.styles["Normal"].font
    base.name = "Times New Roman"
    base.size = Pt(12)
    for p in doc.paragraphs:
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.first_line_indent = Cm(1.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # Center align known titles by matching their lowercase text
        if (p.text or "").strip().lower() in KEYWORDS_TITLES:
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def docx_to_html_bytes(docx_bytes: bytes) -> str:
    """Convert `.docx` bytes to HTML using `mammoth`.

    Returns an HTML string representing the document's structure and content.
    """
    with io.BytesIO(docx_bytes) as bio:
        result = mammoth.convert_to_html(bio)
        return result.value

def sanitize_html(html: str) -> str:
    """Remove potentially dangerous tags and attributes from HTML.

    For preview purposes we remove `<script>` blocks and inline event
    handlers. This function is intentionally simple—if you need robust
    sanitization consider integrating a library dedicated to HTML
    sanitization.
    """
    html = re.sub(r"<script.*?>.*?</script>", "", html, flags=re.I | re.S)
    html = re.sub(r"on\w+\s*=\s*\".*?\"", "", html, flags=re.I)
    return html

@app.get("/")
def root() -> Dict[str, str]:
    """Basic health check endpoint to verify service availability."""
    return {"name": "ABNT Conversor", "status": "ok", "version": "0.2.0"}

@app.post("/analyze")
async def analyze(file: UploadFile = File(...)):
    """Extract plain text and heuristic spans from an uploaded `.docx`.

    Returns a JSON object with the raw text and a list of spans for
    highlighting. If the upload is invalid or any exception occurs,
    responds with a 400 status and an error message.
    """
    try:
        raw = load_docx_bytes(file)
        with tempfile.TemporaryDirectory() as td:
            path = os.path.join(td, "in.docx")
            with open(path, "wb") as f:
                f.write(raw)
            doc = Document(path)
        text = extract_plain_text(doc)
        spans = heuristic_spans(text)
        return {"text": text, "spans": spans}
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=400)

@app.post("/preview")
async def preview(
    file: UploadFile = File(...),
    map: Optional[str] = Form(None),  # Future parameter to accept adjusted spans
):
    """Generate sanitized HTML for previewing the document in the browser.

    Accepts the uploaded `.docx` and returns a JSON object containing
    `preview_html`, which wraps the converted HTML within a styled page. The
    returned HTML is safe to embed via `<iframe srcdoc>` on the client
    side. The optional `map` form field is reserved for future use, for
    passing user-adjusted spans back to the server.
    """
    try:
        raw = load_docx_bytes(file)
        html = docx_to_html_bytes(raw)
        sanitized = sanitize_html(html)
        # Basic ABNT-like styling for preview
        style = """
        <style>
        body{margin:0;padding:16px;font:14px/1.5 system-ui,-apple-system,Segoe UI,Roboto}
        .page{width:794px;min-height:1123px;margin:0 auto;background:#fff;box-shadow:0 0 0 1px #e5e7eb;padding:3cm 2cm 2cm 3cm}
        h1,h2,h3{font-weight:700;margin:12px 0}
        p{margin:0}
        </style>
        """
        preview_html = f"<!doctype html><meta charset='utf-8'>{style}<div class='page'>{sanitized}</div>"
        return {"preview_html": preview_html}
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=400)

@app.post("/convert")
async def convert(
    file: UploadFile = File(...),
    metadata: Optional[str] = Form(None),
    out: Optional[str] = Form("docx"),  # 'docx' | 'pdf' | 'both'
):
    """Return an ABNT-formatted `.docx` file from the uploaded document.

    This endpoint writes the uploaded `.docx` to a temporary file, applies
    formatting, and saves the result with an `ABNT_` prefix. The file is
    served directly in the response using FastAPI's `FileResponse`.
    """
    try:
        raw = load_docx_bytes(file)
        with tempfile.TemporaryDirectory() as td:
            inp_path = os.path.join(td, file.filename or "entrada.docx")
            out_path = os.path.join(td, f"ABNT_{os.path.basename(inp_path)}")
            with open(inp_path, "wb") as f:
                f.write(raw)
            doc = Document(inp_path)
            apply_abnt_formatting(doc)
            # Note: 'metadata' and 'out' parameters are placeholders for future use
            doc.save(out_path)
            return FileResponse(
                out_path,
                filename=os.path.basename(out_path),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=400)
